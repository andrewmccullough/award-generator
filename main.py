#####################
### CONFIGURATION ###
#####################

## Hello. If you are unfamiliar with Python, this file may look intimidating to you. While my comments throughout the code should help you to understand what it is doing, you don't have to worry about most of it.

## Please take a look at the Word document included in the template-filler repository named "template.docx". In it, you will find a table with the portions of the template that are to be modified by this script.

## Each row of this table has a number to identify it, starting at 0 and incrementing upwards. The first row == 0, the second row == 1, etc. You can change which row the script should edit by modifying the variables below this comment.

delegationRow = 1
awardRow = 3
committeeRow = 5

## If you want to modify the template or the script more than these three variables will allow, you should familiarize yourself with the Python package "python-docx" and with the contents of this script.

import os, sys, textwrap

#########################
### WINDOW DIMENSIONS ###
#########################

def window ():

    '''
        Determines the height and width of the Terminal window.
    '''

    env = os.environ
    def ioctl_GWINSZ(fd):
        try:
            import fcntl, termios, struct, os
            cr = struct.unpack('hh', fcntl.ioctl(fd, termios.TIOCGWINSZ,
        '1234'))
        except:
            return
        return cr
    cr = ioctl_GWINSZ(0) or ioctl_GWINSZ(1) or ioctl_GWINSZ(2)
    if not cr:
        try:
            fd = os.open(os.ctermid(), os.O_RDONLY)
            cr = ioctl_GWINSZ(fd)
            os.close(fd)
        except:
            pass
    if not cr:
        cr = (env.get('LINES', 25), env.get('COLUMNS', 80))

    return cr[1], cr[0]

width, height = window ()

#################
### FUNCTIONS ###
#################

def printMessage(message):

    '''
        Prints long messages to the Terminal without splitting words across two lines.
    '''

    message = message.split()
    lines = []

    line = ""

    while len(message) > 0:

        if len(message[0]) >= width:
            line = line + message[0] + " "
            message.pop(0)
            lines.append(line)
            line = ""
        else:
            if len(line) + len(message[0]) < width:

                # If length of current line plus the next word of the message is less than the width of the Terminal, add the next word to the current line.
                line = line + message[0] + " "
                message.pop(0)

                if len(message) == 0:
                    # If the word that was just appended to the current line was the last word in the message, append the current line to the list of lines.
                    lines.append(line)

            else:

                # Otherwise, append the current line to the list of lines and start a new line.
                lines.append(line)
                line = ""


    for l in lines:
        print(l)

def divider():

    '''
        Prints a row of equals signs across the Terminal as a separator.
    '''

    print("=" * width)

def goodbye(message = None):

    '''
        Ends script with "Goodbye." and a row of equals signs.
    '''

    if message is not None:
        printMessage(message)
        divider()

    printMessage("Goodbye.")
    divider()
    sys.exit()

#############################
### EXTERNAL DEPENDENCIES ###
#############################

attempted = False

'''
    Attempts to import "python-docx". If it is not installed, offers to install it if the script has not already tried to do so.
'''

try:
    import docx
except ImportError:
    # On unsuccessful import...
    printMessage("You don't have the \"python-docx\" package installed, which is required for this script. You can use PIP to install it, using the command \"pip3 install python-docx\".")

    if not attempted:
        # If the script has not attempted to install the dependency already...
        response = None
        while response == None:
            print("Would you like us to try to install it?")

            divider()
            response = input(" $ ").lower()
            divider()

            if response in ["yes", "y"]:
                os.system("pip3 install python-docx")
                os.system("clear")
                attempted = True
            elif response in ["no", "n"]:
                goodbye()
            elif response == "":
                goodbye()
            else:
                printMessage("Please enter yes or no.")
                response = None
    else:
        # If the script has already attempted to install the dependency...
        goodbye("We attempted to install the \"python-docx\" package but were unable to do so. Please run the command yourself.")

attempted = False

'''
    Attempts to import "py-xlsx". If it is not installed, offers to install it if the script has not already tried to do so.
'''

try:
    import xlsx
except ImportError:
    # On unsuccessful import...
    printMessage("You don't have the \"py-xlsx\" package installed, which is required for this script. You can use PIP to install it, using the command \"pip3 install py-xlsx\".")

    if not attempted:
        # If the script has not attempted to install the dependency already...
        response = None
        while response == None:
            printMessage("Would you like us to try to install it?")

            divider()
            response = input(" $ ").lower()
            divider()

            if response in ["yes", "y"]:
                os.system("pip3 install py-xlsx")
                os.system("clear")
                attempted = True
            elif response in ["no", "n"]:
                goodbye()
            elif response == "":
                goodbye()
            else:
                printMessage("Please enter yes or no.")
                response = None
    else:
        # If the script has already attempted to install the dependency...
        printMessage("We attempted to install the \"py-xlsx\" package but were unable to do so. Please run the command yourself.")
        goodbye()

###################
### MAIN SCRIPT ###
###################

os.system("clear")

printMessage("Drag and drop your awards spreadsheet from anywhere on your Mac into the Terminal.")

divider()
awardsFile = input(" $ ").strip().replace("\\", "")
divider()

if awardsFile == "":
    # If the user entered no filename...
    goodbye()

if not os.path.isfile(awardsFile):
    # If the filename entered by the user does not exist...
    goodbye("That file does not exist. Make sure you entered the name correctly and that it is in this directory.")

awardsFileExtension = os.path.splitext(os.path.realpath(awardsFile))[1]

if awardsFileExtension not in [".csv", ".xlsx"]:
    # If the file is not a CSV or XLSX file...
    goodbye("The awards spreadsheet must be in either CSV or XLSX format.")

if not os.path.isfile("~/Desktop/template.docx"):
    # If template.docx is not found...
    goodbye("We can't find the template file for your awards [\"template.docx\"] on your Desktop.")

if awardsFileExtension == ".csv":

    '''
        Processes the awards spreadsheet if it is a CSV file. Relevant data moved to $lines.
    '''

    f = open(awardsFile)

    lines = f.readlines()
    lines = [l.strip().split(",") for l in lines]

    f.close()

elif awardsFileExtension == ".xlsx":

    '''
        Processes the awards spreadsheet if it is an Excel file. Relevant data moved to $lines.
    '''

    book = xlsx.Workbook(awardsFile)

    sheetName = False
    for sheet in book:
        if not sheetName:
            sheetName = sheet.name

    sheet = book[sheetName]
    rows = sheet.rows()

    lines = []

    for row in rows:
        line = [c.value for c in rows[row]]
        lines.append(line)

l = lines[0]

if l[0].lower() in ["committee", "position", "delegate", "delegation", "award"] or l[1].lower() in ["committee", "position", "delegate", "delegation", "award"] or l[2].lower() in ["committee", "position", "delegate", "delegation", "award"]:

    '''
        Checks if the first line of awards data could be a header.
    '''

    response = None
    while response is None:
        printMessage("It looks like the first row of your awards spreadsheet may be a header. Should we make an award for the first row?")

        divider()
        response = input(" $ ").lower()
        divider()

        if response in ["yes", "y"]:
            # If user would still like an award made for the first line, even though it appears the first line is a header...
            pass
        elif response in ["no", "n"]:
            # If the user would not like an award made for the first line, as it appears the first line is a header...
            lines.pop(0)
        elif response == "":
            goodbye()
        else:
            printMessage("Please enter yes or no.")
            response = None

response = None
while response is None:

    '''
        Confirms that the data is in the proper order.
    '''

    printMessage("This script needs your awards spreadsheet to be ordered \"committee\", \"award\", \"delegation\". Is it in this order?")

    divider()
    response = input(" $ ").lower()
    divider()

    if response in ["yes", "y"]:
        pass
    elif response in ["no", "n"]:
        goodbye("Please put it in this order.")
    elif response == "":
        goodbye()
    else:
        printMessage("Please enter yes or no.")
        response = None

if not os.path.isdir("exports"):
    # Makes an "exports" directory if one does not already exist.
    os.makedirs("exports")

i = 0 # counter for number of awards made

try:
    for l in lines:
        committee = l[0]
        award = l[1]
        delegation = l[2]

        doc = docx.Document("template.docx")

        fields = doc.tables[0]

        fields.rows[delegationRow].cells[0].paragraphs[0].text = delegation
        fields.rows[awardRow].cells[0].paragraphs[0].text = award
        fields.rows[committeeRow].cells[0].paragraphs[0].text = committee

        filename = "__".join([committee, award, delegation])
        filename = filename + ".docx"

        doc.save("~/Desktop/exports/" + filename)

        i = i + 1
except:
    goodbye("The script encountered a problem.")

printMessage(str(i) + " awards were successfully created.")
printMessage("They are available in " + os.path.realpath("~/Desktop/exports/") + ".")
divider()
